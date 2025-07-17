"""
Utility functions for document processing and model management
"""

import os
import re
import string
import joblib
import logging
from typing import Tuple, Optional

import fitz  # PyMuPDF
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.ensemble import RandomForestClassifier
import numpy as np

# Configure logging
logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)

# Initialize stemmer and stopwords
stemmer = PorterStemmer()
stop_words = set(stopwords.words('english'))

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using PyMuPDF
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as string
    """
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        doc.close()
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise Exception(f"Failed to extract text from PDF: {e}")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as string
    """
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise Exception(f"Failed to extract text from DOCX: {e}")

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from file based on extension
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text as string
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def preprocess_text(text: str) -> str:
    """
    Preprocess text for classification
    
    Args:
        text: Raw text to preprocess
        
    Returns:
        Preprocessed text
    """
    try:
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters and digits
        text = re.sub(r'[^a-zA-Z\s]', '', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Tokenize
        tokens = word_tokenize(text)
        
        # Remove stopwords and apply stemming
        processed_tokens = []
        for token in tokens:
            if token not in stop_words and len(token) > 2:
                stemmed_token = stemmer.stem(token)
                processed_tokens.append(stemmed_token)
        
        return ' '.join(processed_tokens)
    
    except Exception as e:
        logger.error(f"Error preprocessing text: {e}")
        return text  # Return original text if preprocessing fails

def create_stub_model_and_vectorizer() -> Tuple[RandomForestClassifier, TfidfVectorizer]:
    """
    Create a stub model and vectorizer for testing purposes
    This would be replaced with a properly trained model in production
    
    Returns:
        Tuple of (model, vectorizer)
    """
    logger.info("Creating stub model and vectorizer...")
    
    # Sample training data for each category
    sample_data = {
        'Legal': [
            'contract agreement terms conditions legal binding party',
            'lawsuit litigation court case legal proceeding attorney',
            'copyright patent trademark intellectual property legal rights',
            'compliance regulation law statute legal requirement',
            'liability negligence legal responsibility damages claim'
        ],
        'HR': [
            'employee personnel human resources staff management',
            'hiring recruitment job position candidate interview',
            'performance evaluation review feedback employee development',
            'benefits compensation salary payroll employee package',
            'training development skill employee education program'
        ],
        'Finance': [
            'budget financial accounting revenue expense profit loss',
            'investment portfolio asset liability financial statement',
            'tax taxation financial reporting audit accounting',
            'cash flow financial analysis business finance',
            'loan credit debt financial institution banking'
        ],
        'Medical': [
            'patient medical health treatment diagnosis doctor',
            'medication prescription drug therapy medical treatment',
            'hospital clinic medical facility healthcare service',
            'surgery operation medical procedure patient care',
            'medical record health information patient data'
        ],
        'Technical': [
            'software development programming code technical specification',
            'system architecture database server technical infrastructure',
            'api documentation technical manual user guide',
            'bug fix technical issue software maintenance',
            'technical support troubleshooting system configuration'
        ]
    }
    
    # Prepare training data
    texts = []
    labels = []
    
    for category_idx, (category, samples) in enumerate(sample_data.items()):
        for sample in samples:
            texts.append(sample)
            labels.append(category_idx)
    
    # Create and fit vectorizer
    vectorizer = TfidfVectorizer(
        max_features=5000,
        ngram_range=(1, 2),
        min_df=1,
        max_df=0.95,
        stop_words='english'
    )
    
    X = vectorizer.fit_transform(texts)
    
    # Create and train model
    model = RandomForestClassifier(
        n_estimators=100,
        random_state=42,
        max_depth=20,
        min_samples_split=5,
        min_samples_leaf=2
    )
    
    model.fit(X, labels)
    
    logger.info("Stub model and vectorizer created successfully")
    return model, vectorizer

def save_model_and_vectorizer(model: RandomForestClassifier, vectorizer: TfidfVectorizer, model_dir: str = 'model') -> None:
    """
    Save model and vectorizer to disk
    
    Args:
        model: Trained classifier model
        vectorizer: Fitted TfidfVectorizer
        model_dir: Directory to save models
    """
    os.makedirs(model_dir, exist_ok=True)
    
    model_path = os.path.join(model_dir, 'classifier_model.joblib')
    vectorizer_path = os.path.join(model_dir, 'vectorizer.joblib')
    
    joblib.dump(model, model_path)
    joblib.dump(vectorizer, vectorizer_path)
    
    logger.info(f"Model saved to {model_path}")
    logger.info(f"Vectorizer saved to {vectorizer_path}")

def load_model_and_vectorizer(model_dir: str = 'model') -> Tuple[RandomForestClassifier, TfidfVectorizer]:
    """
    Load model and vectorizer from disk, create stub if not found
    
    Args:
        model_dir: Directory containing saved models
        
    Returns:
        Tuple of (model, vectorizer)
    """
    model_path = os.path.join(model_dir, 'classifier_model.joblib')
    vectorizer_path = os.path.join(model_dir, 'vectorizer.joblib')
    
    try:
        if os.path.exists(model_path) and os.path.exists(vectorizer_path):
            logger.info("Loading existing model and vectorizer...")
            model = joblib.load(model_path)
            vectorizer = joblib.load(vectorizer_path)
            logger.info("Model and vectorizer loaded successfully")
            return model, vectorizer
        else:
            logger.info("Model files not found, creating stub model...")
            model, vectorizer = create_stub_model_and_vectorizer()
            save_model_and_vectorizer(model, vectorizer, model_dir)
            return model, vectorizer
    
    except Exception as e:
        logger.error(f"Error loading model: {e}")
        logger.info("Creating stub model as fallback...")
        model, vectorizer = create_stub_model_and_vectorizer()
        save_model_and_vectorizer(model, vectorizer, model_dir)
        return model, vectorizer

def validate_model_performance(model: RandomForestClassifier, vectorizer: TfidfVectorizer) -> dict:
    """
    Validate model performance with test samples
    
    Args:
        model: Trained classifier model
        vectorizer: Fitted TfidfVectorizer
        
    Returns:
        Dictionary with performance metrics
    """
    test_samples = [
        ("This employment contract specifies the terms and conditions of employment", "Legal"),
        ("Employee performance review and feedback session scheduled", "HR"),
        ("Quarterly financial report showing revenue and expenses", "Finance"),
        ("Patient medical history and treatment recommendations", "Medical"),
        ("Software API documentation and technical specifications", "Technical")
    ]
    
    correct_predictions = 0
    total_predictions = len(test_samples)
    categories = ['Legal', 'HR', 'Finance', 'Medical', 'Technical']
    
    for text, expected_category in test_samples:
        processed_text = preprocess_text(text)
        text_vector = vectorizer.transform([processed_text])
        prediction = model.predict(text_vector)[0]
        predicted_category = categories[prediction]
        
        if predicted_category == expected_category:
            correct_predictions += 1
    
    accuracy = correct_predictions / total_predictions
    
    return {
        'accuracy': accuracy,
        'correct_predictions': correct_predictions,
        'total_predictions': total_predictions,
        'test_samples': test_samples
    }